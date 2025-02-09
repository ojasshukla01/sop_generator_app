from docx import Document

def generate_sop(user_details: dict, sop_content: dict) -> str:
    """Generate an SOP document in .docx format."""
    doc = Document()
    doc.add_heading("Statement of Purpose", 0)

    doc.add_paragraph(f"Name: {user_details['name']}")
    doc.add_paragraph("Introduction:")
    doc.add_paragraph(sop_content['introduction'])

    # Add other sections
    doc.add_paragraph("Academic Background:")
    doc.add_paragraph(sop_content['academic_background'])

    doc.add_paragraph("Course Details:")
    doc.add_paragraph(sop_content['course_details'])

    doc.save("output/sop.docx")  # Save the SOP as a .docx file
    return "output/sop.docx"
